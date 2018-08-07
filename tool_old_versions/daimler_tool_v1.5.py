# coding=utf-8
from docx import Document
import shutil
import os
import stat
import datetime
import pyexcel
import openpyxl
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Side, PatternFill
import ctypes  # An included library with Python install.

print"Modules imported."

def message_box(title, message):
    ctypes.windll.user32.MessageBoxA(None, message, title, 0)


files = os.listdir(os.curdir)
indus_folder_name = os.getcwd().split("\\")[-1]

error_count = 0
key_position = 0

selected_indus_number = 358

source_xl = 'EMS_Meeting_Minutes_2018.xls'
destination_xl = 'Follow-up_CR_CN_ECU_Daimler.xlsx'
destination_sheet_save_as = 'Follow-up_CR_CN_ECU_Daimler.xlsx'
destination_sheet_name = 'DAIMLER 651'

if not os.path.isfile(source_xl):
    message_box("Error", 'Missing' + source_xl + 'file.')
elif not os.path.isfile(destination_xl):
    message_box("Error", 'Missing' + destination_xl + 'file.')
else:
    source_sheet = pyexcel.get_sheet(file_name=source_xl)
    source_sheet.name_columns_by_row(0)
    print source_xl + " found and opened."

    destination_book = openpyxl.load_workbook(destination_xl)
    destination_sheet = destination_book[destination_sheet_name]
    print destination_xl + " found and opened."

    my_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'),
                       bottom=Side(style='thin'))
    blank_row_color = PatternFill(start_color='a6a6a6', end_color='a6a6a6', fill_type='solid')
    delivery_color = PatternFill(start_color='fde9d9', end_color='a6a6a6', fill_type='solid')
    print "Tables styles created."

    item_quantity = 0
    values_to_xl = []
    blank_row = (' ', ' ', ' ', ' ', ' ')

    os.chmod('Follow-up_CR_CN_ECU_Daimler.xlsx', stat.S_IWRITE)
    print "File read and write permission granted."

    num_of_docx = len([name for name in os.listdir('.') if name.endswith('.docx')])
    num_of_s19 = len([name for name in os.listdir('.') if name.endswith('.s19')])
    print "Number of files checked."

    no_file_in_root_folder = num_of_docx == 0 and num_of_s19 == 0

    if indus_folder_name.startswith("ME"):
        selected_indus_number = indus_folder_name
    elif indus_folder_name.startswith("Indus"):
        selected_indus_number = indus_folder_name.split("Indus")[-1]
    print "Indus number read from folder name."

    found_in_xl = 0
    append_before = 0

    # find indus_number in sheet and get some values(codes_with_delivery, ecu_type, sw_verson, cr_number, cr_creation)
    for row in source_sheet:
        if row[3] == int(selected_indus_number):
            found_in_xl = 1
            codes_with_delivery = row[2]
            ecu_type = row[4][3:]
            sw_version = row[6]
            cr_number = row[8]
            cr_creation = row[10]
            break
    print "Data collected belong to selected indus number."

    # check for appended before
    for i in range(1, destination_sheet.max_row):
        if destination_sheet.cell(row=i, column=5).value == sw_version and destination_sheet.cell(row=i, column=6).value == int(selected_indus_number):
            append_before = 1
            break
    print "Checked whether values ​​were added before."

    if found_in_xl == 0:
        message_box("Error", "Indus number could'nt find in Excel Sheet.")
    elif append_before == 1:
        message_box("Error", "Appended Before")
    else:
        if codes_with_delivery == "" or ecu_type == "" or sw_version == "" or cr_number == "" or cr_creation == "":
            message_box("Error", "This indus number has missing values.")
        else:
            if no_file_in_root_folder:

                folder_list = filter(lambda x: os.path.isdir(x), os.listdir('.'))

                for folders in folder_list:
                    files = os.listdir(folders)

                    num_of_docx = len([name for name in os.listdir(folders) if name.endswith('.docx')])
                    num_of_s19 = len([name for name in os.listdir(folders) if name.endswith('.s19')])

                    if num_of_s19 == num_of_docx and num_of_s19 == 1:
                        for f in files:
                            if f.endswith('.docx'):
                                file_name = f

                                document = Document(folders + '/' + file_name)
                                tables = document.tables

                                if not tables[0].rows[31].cells[2].text:
                                    error_count += 1

                                    if not tables[0].rows[32].cells[2].text:

                                        if not tables[0].rows[33].cells[2].text:

                                            if not tables[0].rows[34].cells[2].text:
                                                key_position = 0
                                                with open("Error-Logs.txt", "a") as text_file:
                                                    text_file.write(
                                                        str(error_count) + "-  " + file_name + " file has no code.\n")
                                                    text_file.write(
                                                        "-----------------------------------------------------------------------------------\n")
                                            else:
                                                unique_key = tables[0].rows[34].cells[2].text
                                                key_position = 4
                                        else:
                                            unique_key = tables[0].rows[33].cells[2].text
                                            key_position = 3
                                    else:
                                        unique_key = tables[0].rows[32].cells[2].text
                                        key_position = 2

                                    if key_position != 0:
                                        with open("Error-Logs.txt", "a") as text_file:
                                            text_file.write(
                                                str(error_count) + "- " + file_name + " 's code is in " + str(
                                                    key_position) + ". line.\n")
                                            text_file.write(
                                                "------------------------------------------------------------------------------ \n")

                                else:
                                    unique_key = tables[0].rows[31].cells[2].text
                                    key_position = 1

                                if key_position != 0:
                                    unique_key_file_name = unique_key + '.s19'
                                    sw_update_code = tables[0].rows[6].cells[3].text

                                    daimler_part_number = file_name.split("_")[1]
                                    file_name_first = "-".join(unique_key.split("-")[3:-1])
                                    file_name_sec = unique_key.split("-")[-1].split("_")[0]
                                    folder_name = file_name_first + '-' + file_name_sec + '-' + daimler_part_number

                                    title = "-".join(unique_key.split("-")[0:3]) + '-' + daimler_part_number
                                    ref_daimler = daimler_part_number
                                    item_quantity += 1
                                    comment = unique_key_file_name

                                    sw_title = "-".join(unique_key.split("-")[0:3]) + '-' + sw_update_code.replace(" ",
                                                                                                                   "")
                                    sw_ref_daimler = sw_update_code.replace(" ", "")
                                    sw_item_quantity = 1
                                    sw_comment = "CODE"

                                    values_to_xl.append(
                                        (cr_number, cr_creation, '', '', sw_version, int(selected_indus_number),
                                         ecu_type,
                                         title, ref_daimler, '', item_quantity, '', '', '', '', '', '', '', '', '',
                                         comment))
                    else:
                        message_box("Error", 'Problem founded in folder. Please check' + folders + 'named folder.')

            else:
                for f in files:
                    if f.endswith('.docx'):
                        file_name = f

                        document = Document(file_name)
                        tables = document.tables

                        if not tables[0].rows[31].cells[2].text:
                            error_count += 1

                            if not tables[0].rows[32].cells[2].text:

                                if not tables[0].rows[33].cells[2].text:

                                    if not tables[0].rows[34].cells[2].text:
                                        key_position = 0
                                        with open("Error-Logs.txt", "a") as text_file:
                                            text_file.write(
                                                str(error_count) + "-  " + file_name + " file has no code.\n")
                                            text_file.write(
                                                "-----------------------------------------------------------------------------------\n")
                                    else:
                                        unique_key = tables[0].rows[34].cells[2].text
                                        key_position = 4
                                else:
                                    unique_key = tables[0].rows[33].cells[2].text
                                    key_position = 3
                            else:
                                unique_key = tables[0].rows[32].cells[2].text
                                key_position = 2

                            if key_position != 0:
                                with open("Error-Logs.txt", "a") as text_file:
                                    text_file.write(
                                        str(error_count) + "- " + file_name + " 's code is in " + str(
                                            key_position) + ". line.\n")
                                    text_file.write(
                                        "------------------------------------------------------------------------------ \n")


                        else:
                            unique_key = tables[0].rows[31].cells[2].text
                            key_position = 1

                        if key_position != 0:

                            unique_key_file_name = unique_key + '.s19'
                            sw_update_code = tables[0].rows[6].cells[3].text

                            daimler_part_number = file_name.split("_")[1]
                            file_name_first = "-".join(unique_key.split("-")[3:-1])
                            file_name_sec = unique_key.split("-")[-1].split("_")[0]
                            folder_name = file_name_first + '-' + file_name_sec + '-' + daimler_part_number

                            title = "-".join(unique_key.split("-")[0:3]) + '-' + daimler_part_number
                            ref_daimler = daimler_part_number
                            item_quantity += 1
                            comment = unique_key_file_name

                            sw_title = "-".join(unique_key.split("-")[0:3]) + '-' + sw_update_code.replace(" ", "")
                            sw_ref_daimler = sw_update_code.replace(" ", "")
                            sw_item_quantity = 1
                            sw_comment = "CODE"

                            # <---- Create directory and moving process ----->
                            if os.path.exists(unique_key_file_name):
                                if not os.path.exists(folder_name):
                                    os.makedirs(folder_name)

                                    shutil.move(unique_key_file_name, folder_name)
                                    shutil.move(file_name, folder_name)

                                    values_to_xl.append(
                                        (cr_number, cr_creation, '', '', sw_version, int(selected_indus_number),
                                         ecu_type,
                                         title, ref_daimler, '', item_quantity, '', '', '', '', '', '', '', '', '',
                                         comment))
                            # <---- Create directory and moving process ----->

                            else:
                                error_count += 1
                                if not os.path.exists("Error-Logs.txt"):
                                    with open("Error-Logs.txt", "w") as text_file:
                                        text_file.write(
                                            str(error_count) + "-  " + file_name + " 's code dont match with\n    "
                                            + unique_key_file_name + "\n")
                                        text_file.write(
                                            "---------------------------------------------------------------------------\n")
                                else:
                                    with open("Error-Logs.txt", "a") as text_file:
                                        text_file.write(
                                            str(error_count) + "-  " + file_name + " 's code dont match with\n    "
                                            + unique_key_file_name + "\n")
                                        text_file.write(
                                            "------------------------------------------------------------------------- \n")

            if codes_with_delivery == 1:
                values_to_xl.append(
                    (cr_number, cr_creation, '', '', sw_version, int(selected_indus_number), ecu_type, sw_title,
                     sw_ref_daimler, '', sw_item_quantity, '', '', '', '', '', '', '', '', '', sw_comment))

            last_row = destination_sheet.max_row
            i = 1
            # <-----List Elements to Excel and Styling--------->

            for row in values_to_xl:
                destination_sheet.append(row)
                for col in range(1, 12):
                    destination_sheet.cell(row=last_row + i, column=col).alignment = openpyxl.styles.Alignment(
                        horizontal='center',
                        vertical='center')
                    destination_sheet.cell(row=last_row + i, column=col).border = my_border
                    if codes_with_delivery == 1:
                        destination_sheet.cell(row=last_row + i, column=col).fill = delivery_color

                destination_sheet.cell(row=last_row + i, column=2).number_format = 'DD.MM.YYYY'

                destination_sheet.cell(row=last_row + i, column=21).alignment = openpyxl.styles.Alignment(
                    horizontal='center',
                    vertical='center')
                destination_sheet.cell(row=last_row + i, column=21).border = my_border
                if codes_with_delivery == 1:
                    destination_sheet.cell(row=last_row + i, column=21).fill = delivery_color

                i += 1

            # <-----List Elements to Excel and Styling--------->

            # <----------------Blank Row Part------------------>

            destination_sheet.append(blank_row)

            for col in range(1, 12):
                destination_sheet.cell(row=last_row + i, column=col).fill = blank_row_color
                destination_sheet.cell(row=last_row + i, column=col).border = my_border

            destination_sheet.cell(row=last_row + i, column=21).fill = blank_row_color
            destination_sheet.cell(row=last_row + i, column=21).border = my_border

            # <----------------End of Blank Row Part----------->

            # <----------------Save File----------------------->

            destination_book.save(destination_sheet_save_as)

        message_box('Complete', 'Process Completed.')


with open("Error-Logs.txt", "a") as text_file:
    text_file.write("Total Error:" + str(error_count) + "   Time:" + str(datetime.datetime.now().time()) + "\n")
    text_file.write("-----------------------------------------------------------------------------------------------\n")

# <--------------End of Save File------------------->
